#!/usr/bin/env python3
"""
Word Document Translation Extraction Script

Parses Word documents to extract formatted translation text delimited by
bold Unicode quotation marks (U+201C and U+201D) and stores them in PostgreSQL.

Usage:
    python parse_word_translations.py
    python parse_word_translations.py --directory "C:\\path\\to\\docs"
    python parse_word_translations.py --dry-run --verbose
"""

import os
import re
import sys
import argparse
import logging
import tempfile
from pathlib import Path
from typing import List, Dict, Optional
from enum import Enum

try:
    from docx import Document
    from docx.text.run import Run
    from docx.table import Table
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    import psycopg2
    from psycopg2 import sql
    from psycopg2.extensions import ISOLATION_LEVEL_AUTOCOMMIT
except ImportError:
    print("ERROR: psycopg2 not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    from dotenv import load_dotenv
except ImportError:
    print("ERROR: python-dotenv not installed. Run: pip install -r requirements.txt")
    sys.exit(1)

try:
    import win32com.client  # noqa: F401 - verified available for subprocess COM
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

# Constants
LEFT_QUOTE = "\u201C"  # "
RIGHT_QUOTE = "\u201D"  # "
DEFAULT_DIRECTORY = r"C:\users\joe\work\dev\yada\docs"


class ExtractionState(Enum):
    """State machine states for translation extraction."""
    SEARCHING = 1
    EXTRACTING = 2
    FOUND_QUOTE = 3
    COMPLETE = 4


def setup_logging(verbose: bool = False) -> None:
    """Configure logging based on verbosity level."""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )


def get_db_connection():
    """
    Create PostgreSQL database connection using environment variables.

    Environment Variables:
        POSTGRES_HOST (default: localhost)
        POSTGRES_PORT (default: 5432)
        POSTGRES_USER (default: postgres)
        POSTGRES_PASSWORD (required)
        POSTGRES_DB (default: yada)

    Returns:
        psycopg2 connection object

    Raises:
        SystemExit if POSTGRES_PASSWORD not set or connection fails
    """
    load_dotenv()

    host = os.getenv('POSTGRES_HOST', 'localhost')
    port = os.getenv('POSTGRES_PORT', '5432')
    user = os.getenv('POSTGRES_USER', 'postgres')
    password = os.getenv('POSTGRES_PASSWORD')

    if not password:
        logging.error("POSTGRES_PASSWORD environment variable not set")
        logging.error("Set it with: set POSTGRES_PASSWORD=your_password (Windows CMD)")
        logging.error("            $env:POSTGRES_PASSWORD='your_password' (Windows PowerShell)")
        sys.exit(1)

    try:
        # First connect to postgres database to check if yada database exists
        conn = psycopg2.connect(
            host=host,
            port=port,
            user=user,
            password=password,
            database='postgres'
        )
        logging.info(f"Connected to PostgreSQL at {host}:{port}")
        return conn
    except psycopg2.Error as e:
        logging.error(f"Failed to connect to PostgreSQL: {e}")
        sys.exit(1)


def init_database(conn) -> None:
    """
    Create database 'yada' if it doesn't exist, then create translation table.

    Args:
        conn: PostgreSQL connection object (connected to postgres database)
    """
    db_name = os.getenv('POSTGRES_DB', 'yada')

    try:
        # Set connection to autocommit mode for database creation
        conn.set_isolation_level(ISOLATION_LEVEL_AUTOCOMMIT)
        cursor = conn.cursor()

        # Check if database exists
        cursor.execute("SELECT 1 FROM pg_database WHERE datname = %s", (db_name,))
        exists = cursor.fetchone()

        if not exists:
            cursor.execute(sql.SQL("CREATE DATABASE {}").format(
                sql.Identifier(db_name)
            ))
            logging.info(f"Created database '{db_name}'")
        else:
            logging.info(f"Database '{db_name}' already exists")

        cursor.close()
        conn.close()

        # Reconnect to the yada database
        host = os.getenv('POSTGRES_HOST', 'localhost')
        port = os.getenv('POSTGRES_PORT', '5432')
        user = os.getenv('POSTGRES_USER', 'postgres')
        password = os.getenv('POSTGRES_PASSWORD')

        conn = psycopg2.connect(
            host=host,
            port=port,
            user=user,
            password=password,
            database=db_name
        )

        cursor = conn.cursor()

        # Create translation table if it doesn't exist
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS translation (
                translation_id SERIAL PRIMARY KEY,
                translation_book VARCHAR(255) NOT NULL,
                translation_page INTEGER,
                translation_text_word TEXT NOT NULL,
                translation_cite VARCHAR(500),
                translation_cite_hebrew VARCHAR(255),
                translation_cite_common VARCHAR(255),
                translation_cite_chapter INTEGER,
                translation_cite_verse INTEGER,
                translation_cite_verse_end SMALLINT,
                translation_cite_note VARCHAR(500),
                translation_cite_book_id INTEGER,
                translation_created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                yy_volume_id INTEGER
            )
        """)

        # Create cite table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS cite (
                id SERIAL PRIMARY KEY,
                label VARCHAR(500) NOT NULL UNIQUE,
                sort SMALLINT NOT NULL DEFAULT 0
            )
        """)

        # Create yy_series table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS yy_series (
                yy_series_id SERIAL PRIMARY KEY,
                yy_series_number SMALLINT,
                yy_series_label VARCHAR(250),
                yy_series_sort SMALLINT DEFAULT 0
            )
        """)

        # Create yy_volume table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS yy_volume (
                yy_volume_id SERIAL PRIMARY KEY,
                yy_series_id INTEGER REFERENCES yy_series(yy_series_id),
                yy_volume_number SMALLINT,
                yy_volume_label VARCHAR(500),
                yy_volume_flip_code VARCHAR(8),
                yy_volume_pdf VARCHAR(500)
            )
        """)

        # Create indexes
        cursor.execute("""
            CREATE INDEX IF NOT EXISTS idx_book ON translation(translation_book)
        """)
        cursor.execute("""
            CREATE INDEX IF NOT EXISTS idx_page ON translation(translation_page)
        """)

        conn.commit()
        logging.info("Database tables and indexes created/verified")

        cursor.close()
        return conn

    except psycopg2.Error as e:
        logging.error(f"Database initialization failed: {e}")
        sys.exit(1)


def get_page_numbers_for_doc_com(doc_path: str, para_texts: Dict[int, str], timeout: int = 120) -> Dict[int, int]:
    """
    Get page numbers for one document using a VBScript that automates Word.

    Uses text-based Find to locate paragraphs instead of index-based lookup,
    avoiding the mismatch between python-docx and COM paragraph indexing.

    Args:
        doc_path: Absolute path to .docx file
        para_texts: Dict mapping para_index -> first ~150 chars of paragraph text
        timeout: Seconds to wait before giving up

    Returns:
        Dict mapping para_index -> page_number
    """
    import subprocess, tempfile

    if not para_texts:
        return {}

    # Sort by para index so searches progress through the document in order
    sorted_items = sorted(para_texts.items(), key=lambda x: x[0])

    # Write search texts to temp file: para_idx|search_text
    inp = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-16')
    for idx, text in sorted_items:
        # Escape pipe characters in text and trim
        safe_text = text.replace('|', ' ').replace('\r', ' ').replace('\n', ' ')
        inp.write(f"{idx}|{safe_text}\n")
    inp.close()

    out = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
    out.close()

    vbs_content = f'''
Dim word, fso, inputFile, outputFile
Dim line, parts, paraIdx, searchText, pageNum

Set fso = CreateObject("Scripting.FileSystemObject")
Set word = CreateObject("Word.Application")
word.Visible = False
word.DisplayAlerts = 0

Set inputFile = fso.OpenTextFile("{inp.name}", 1, False, -1)
Set outputFile = fso.CreateTextFile("{out.name}", True, False)

WScript.StdErr.WriteLine "Opening document..."
Dim doc
Set doc = word.Documents.Open("{doc_path}", , True)
WScript.StdErr.WriteLine "Opened, paragraphs: " & doc.Paragraphs.Count

Dim rng
Dim lastPos
lastPos = 0

Do While Not inputFile.AtEndOfStream
    line = inputFile.ReadLine()
    If Len(line) > 0 Then
        Dim pipePos
        pipePos = InStr(line, "|")
        If pipePos > 0 Then
            paraIdx = Left(line, pipePos - 1)
            searchText = Mid(line, pipePos + 1)

            If Len(searchText) > 0 Then
                Set rng = doc.Content.Duplicate
                rng.SetRange lastPos, doc.Content.End
                rng.Find.ClearFormatting
                rng.Find.Text = searchText
                rng.Find.Forward = True
                rng.Find.Wrap = 0
                rng.Find.MatchWildcards = False
                rng.Find.MatchCase = True

                On Error Resume Next
                rng.Find.Execute
                If Err.Number = 0 And rng.Find.Found Then
                    pageNum = rng.Information(1)
                    outputFile.WriteLine paraIdx & "|" & pageNum
                    lastPos = rng.Start
                End If
                On Error GoTo 0
            End If
        End If
    End If
Loop

doc.Close False
inputFile.Close
outputFile.Close
word.Quit
WScript.StdErr.WriteLine "Done"
'''

    vbs_file = tempfile.NamedTemporaryFile(mode='w', suffix='.vbs', delete=False, encoding='utf-8')
    vbs_file.write(vbs_content)
    vbs_file.close()

    try:
        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_file.name],
            capture_output=True, text=True, timeout=timeout
        )

        if result.stderr:
            for line in result.stderr.strip().split('\n'):
                if line.strip():
                    logging.debug(f"  VBS: {line.strip()}")

        page_map = {}
        if os.path.exists(out.name):
            with open(out.name, 'r', encoding='utf-8', errors='replace') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('\ufeff'):
                        line = line[1:]
                    parts = line.split('|')
                    if len(parts) >= 2:
                        page_map[int(parts[0])] = int(parts[1])

        return page_map

    except subprocess.TimeoutExpired:
        logging.warning(f"  COM timed out for {Path(doc_path).name} after {timeout}s")
        try:
            subprocess.run(["powershell", "-Command",
                "Stop-Process -Name WINWORD -Force -ErrorAction SilentlyContinue"],
                capture_output=True, timeout=10)
        except:
            pass
        return {}
    except Exception as e:
        logging.warning(f"  COM failed for {Path(doc_path).name}: {e}")
        return {}
    finally:
        for fp in [inp.name, out.name, vbs_file.name]:
            try:
                os.unlink(fp)
            except:
                pass



def build_page_map_from_xml(doc_path: Path) -> Dict[int, int]:
    """
    Build a paragraph-to-page map from XML page breaks and section restarts.

    Fallback method when COM automation is unavailable or times out.
    Uses lastRenderedPageBreak and explicit page breaks from the XML.

    Returns python-docx-compatible indices (top-level paragraphs only),
    but tracks page breaks in ALL paragraphs including those inside tables.

    Args:
        doc_path: Path to .docx file

    Returns:
        Dict mapping python-docx 0-based paragraph index -> estimated page number
    """
    doc = Document(str(doc_path))
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    body = doc.element.body

    # Use findall for both to get consistent Python object identities
    direct_paras = body.findall('w:p', nsmap)
    all_paras = body.findall('.//w:p', nsmap)
    top_level_ids = set(id(p) for p in direct_paras)

    # Find section page number restarts (keyed by absolute para index)
    section_starts = {}
    for sect in body.findall('.//w:sectPr', nsmap):
        pg_num = sect.find('w:pgNumType', nsmap)
        if pg_num is not None:
            start_val = pg_num.get(f'{{{nsmap["w"]}}}start')
            if start_val:
                parent = sect.getparent()
                if parent.tag == f'{{{nsmap["w"]}}}body':
                    # Last sectPr in body - applies to last paragraph
                    pass
                elif parent.tag == f'{{{nsmap["w"]}}}pPr':
                    p_elem = parent.getparent()
                    for pi, pe in enumerate(all_paras):
                        if pe is p_elem:
                            section_starts[pi] = int(start_val)
                            break

    # Build page map tracking ALL paragraphs but only recording top-level ones
    page = 1
    para_to_page = {}
    docx_idx = 0

    for abs_idx, para_elem in enumerate(all_paras):
        # Check for section restart
        if abs_idx in section_starts:
            page = section_starts[abs_idx]

        # Only record page for top-level paragraphs (matching doc.paragraphs)
        if id(para_elem) in top_level_ids:
            para_to_page[docx_idx] = page
            docx_idx += 1

        # Check for page breaks within this paragraph (including table paras)
        # Only count lastRenderedPageBreak (placed by Word on last save);
        # explicit <w:br type="page"> breaks are already reflected in lastRenderedPageBreak
        for br in para_elem.findall('.//w:lastRenderedPageBreak', nsmap):
            page += 1

    return para_to_page


def get_all_page_numbers(doc_para_map: Dict[str, List[int]]) -> Dict[str, Dict[int, int]]:
    """
    Get page numbers for all documents. Tries COM per document, falls back to XML.

    Uses text-based Find in VBScript to locate paragraphs, avoiding mismatches
    between python-docx and COM paragraph indexing.

    Args:
        doc_para_map: Dict mapping absolute doc path -> list of 0-based paragraph indices

    Returns:
        Dict mapping doc path -> {para_index: page_number}
    """
    if not doc_para_map:
        return {}

    total_paras = sum(len(v) for v in doc_para_map.values())
    total_docs = len(doc_para_map)
    logging.info(f"Getting page numbers for {total_docs} documents, {total_paras} paragraphs...")

    all_page_numbers = {}
    com_success = 0
    xml_fallback = 0

    # Sort by file size (smaller first - they work better with COM)
    sorted_docs = sorted(doc_para_map.items(), key=lambda x: Path(x[0]).stat().st_size)

    for doc_path, indices in sorted_docs:
        doc_name = Path(doc_path).name
        file_size_mb = Path(doc_path).stat().st_size / (1024 * 1024)

        # Timeout based on file size: 60s base + 30s per MB
        timeout = int(60 + file_size_mb * 30)

        logging.info(f"  {doc_name} ({file_size_mb:.1f}MB, {len(indices)} paras, timeout {timeout}s)")

        # Get paragraph text for text-based Find (avoids index mapping issues)
        doc = Document(str(doc_path))
        para_texts = {}
        for idx in indices:
            if idx < len(doc.paragraphs):
                raw_text = doc.paragraphs[idx].text.strip()
                # Use first 150 chars as search text (enough for uniqueness)
                snippet = raw_text[:150] if len(raw_text) > 150 else raw_text
                if snippet:
                    para_texts[idx] = snippet

        page_map = get_page_numbers_for_doc_com(doc_path, para_texts, timeout=timeout)

        if page_map:
            all_page_numbers[doc_path] = page_map
            com_success += 1
            logging.info(f"    COM: {len(page_map)}/{len(indices)} page numbers")
        else:
            # Try re-saving via python-docx to create a clean copy (strips problematic elements)
            logging.info(f"    COM failed on original, trying clean copy via python-docx...")
            try:
                clean_dir = tempfile.mkdtemp()
                clean_path = os.path.join(clean_dir, "clean_copy.docx")
                clean_doc = Document(str(doc_path))
                clean_doc.save(clean_path)
                page_map = get_page_numbers_for_doc_com(clean_path, para_texts, timeout=timeout)
                try:
                    os.unlink(clean_path)
                    os.rmdir(clean_dir)
                except:
                    pass
            except Exception as e:
                logging.warning(f"    Clean copy attempt failed: {e}")
                page_map = {}

            if page_map:
                all_page_numbers[doc_path] = page_map
                com_success += 1
                logging.info(f"    COM (clean copy): {len(page_map)}/{len(indices)} page numbers")
            else:
                # Fall back to XML-based page mapping using lastRenderedPageBreak + pgNumType restarts
                logging.info(f"    COM failed on clean copy too, using XML fallback")
                try:
                    full_page_map = build_page_map_from_xml(Path(doc_path))
                    page_map = {idx: full_page_map.get(idx, 1) for idx in indices}
                    all_page_numbers[doc_path] = page_map
                    xml_fallback += 1
                    logging.info(f"    XML: {len(page_map)} page numbers (from lastRenderedPageBreak + pgNumType)")
                except Exception as e:
                    logging.warning(f"    XML fallback also failed: {e}")

    logging.info(f"Page numbers: {com_success} docs via COM, {xml_fallback} docs via XML fallback")
    return all_page_numbers



def replace_unicode_chars(text: str) -> str:
    """
    Replace private use area Unicode characters with their intended letters.

    Args:
        text: Input text

    Returns:
        Text with replacements
    """
    # Replace private use area characters
    text = text.replace('\uf065', 'e')
    text = text.replace('\uf066', 'f')
    text = text.replace('\uf069', 'i')
    return text


def format_run_as_html(run: Run) -> str:
    """
    Convert a python-docx Run object to HTML with formatting and font detection.

    Args:
        run: python-docx Run object

    Returns:
        HTML-formatted string with font spans and format tags
    """
    text = run.text

    # Replace special Unicode characters first
    text = replace_unicode_chars(text)

    # Check for non-default font
    font_name = None
    try:
        if run.font and run.font.name and run.font.name != 'Times New Roman':
            font_name = run.font.name
    except:
        pass

    # Apply font span if non-default
    if font_name:
        text = f'<span class="{font_name}">{text}</span>'

    # Apply formatting in order: underline -> italic -> bold (innermost to outermost)
    if run.underline:
        text = f"<u>{text}</u>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.bold:
        text = f"<b>{text}</b>"

    return text


def extract_cite(text_after_quote: str) -> Optional[str]:
    """
    Extract citation from text following the closing quote.

    Args:
        text_after_quote: Text immediately after the right quote (may span multiple runs)

    Returns:
        Citation text without parentheses, or None if no match
    """
    # Replace special Unicode characters
    text_after_quote = replace_unicode_chars(text_after_quote)

    # Pattern: optional whitespace/nbsp, then parentheses with content
    # More flexible pattern to handle various spacing
    match = re.match(r'^[\s\u00a0]*\(([^)]+)\)', text_after_quote)
    if match:
        cite = match.group(1).strip()
        logging.debug(f"Extracted cite: {cite}")
        return cite
    return None


def parse_cite(cite_text: Optional[str]) -> tuple:
    """
    Split a citation into name, chapter, verse, verse_end, and note components.

    Input: "Mizmowr / Song / Psalm 29:2 - in part"
    Output: ("Mizmowr / Song / Psalm", 29, 2, None, "in part")

    Input: "Yirma'yah / Yah Uplifts / Jeremiah 7:17-18"
    Output: ("Yirma'yah / Yah Uplifts / Jeremiah", 7, 17, 18, None)

    Also handles bare "6:18" -> (None, 6, 18, None, None)

    Args:
        cite_text: Full citation text

    Returns:
        Tuple of (cite_name, cite_chapter, cite_verse, cite_verse_end, cite_note)
    """
    if not cite_text:
        return (None, None, None, None, None)

    # Try "Book Name chapter:verse[-verse_end] [optional note]" pattern first
    match = re.match(r'^(.+?)\s+(\d+):(\d+)(?:-(\d+))?\s*(.*)$', cite_text)
    if match:
        name = match.group(1).strip()
        chapter = int(match.group(2))
        verse = int(match.group(3))
        verse_end = int(match.group(4)) if match.group(4) else None
        note_raw = match.group(5).strip() if match.group(5) else None
        # Strip non-letter characters from beginning and end of note
        note = None
        if note_raw:
            note = re.sub(r'^[^a-zA-Z]+|[^a-zA-Z]+$', '', note_raw)
            if not note:
                note = None
        return (name, chapter, verse, verse_end, note)

    # Handle bare "chapter:verse[-verse_end] [optional note]" without book name
    match = re.match(r'^(\d+):(\d+)(?:-(\d+))?\s*(.*)$', cite_text)
    if match:
        chapter = int(match.group(1))
        verse = int(match.group(2))
        verse_end = int(match.group(3)) if match.group(3) else None
        note_raw = match.group(4).strip() if match.group(4) else None
        note = None
        if note_raw:
            note = re.sub(r'^[^a-zA-Z]+|[^a-zA-Z]+$', '', note_raw)
            if not note:
                note = None
        return (None, chapter, verse, verse_end, note)

    return (cite_text, None, None, None, None)


def split_cite_name(cite_name: Optional[str]) -> tuple:
    """
    Extract Hebrew and common name from a cite name containing slashes.

    Input: "Yirma'yah / Yah Uplifts / Jeremiah"
    Output: ("Yirma'yah", "Jeremiah")

    Input: "Yownah / Jonah"
    Output: ("Yownah", "Jonah")

    Input: "SomeName" (no slashes)
    Output: ("SomeName", None)

    Args:
        cite_name: Full cite name string

    Returns:
        Tuple of (cite_hebrew, cite_common)
    """
    if not cite_name:
        return (None, None)
    if '/' not in cite_name:
        return (cite_name.strip(), None)
    parts = cite_name.split('/')
    cite_hebrew = parts[0].strip()
    cite_common = parts[-1].strip()
    return (cite_hebrew, cite_common)


def consolidate_html(html: str) -> str:
    """
    Consolidate adjacent identical HTML tags into single tags.

    Merges patterns like:
      <span class="X">a</span><span class="X">b</span> -> <span class="X">ab</span>
      </b><b> -> (removed)
      </i><i> -> (removed)
      </u><u> -> (removed)

    Args:
        html: HTML string with formatting tags

    Returns:
        Consolidated HTML string
    """
    # Merge adjacent identical spans (iterative until no more matches)
    prev = None
    while prev != html:
        prev = html
        html = re.sub(
            r'<span class="([^"]+)">([^<]*)</span><span class="\1">',
            r'<span class="\1">\2',
            html
        )

    # Remove empty boundaries for bold, italic, underline
    for tag in ('b', 'i', 'u'):
        html = html.replace(f'</{tag}><{tag}>', '')

    return html



def extract_translations_from_doc(doc_path: Path) -> List[Dict[str, any]]:
    """
    Extract all translations from a Word document.

    Args:
        doc_path: Path to .docx file

    Returns:
        List of dictionaries with keys: book, page, text_word, cite, cite_chapter, cite_verse
    """
    translations = []

    try:
        doc = Document(str(doc_path))
        book_name = doc_path.stem  # Filename without extension

        logging.info(f"Processing document: {book_name}")

        # State machine variables
        state = ExtractionState.SEARCHING
        accumulated_html = []
        start_paragraph_index = 0
        page_number = None
        cite_search_text = ""  # Accumulate text after right quote for cite search

        for para_idx, paragraph in enumerate(doc.paragraphs):
            logging.debug(f"Paragraph {para_idx}: {paragraph.text[:50]}...")

            # Track where this paragraph's HTML contributions start
            para_html_start = len(accumulated_html) if state == ExtractionState.EXTRACTING else None
            # Character position where extraction starts in this paragraph (0 for continuation paras)
            para_extract_start = 0 if state == ExtractionState.EXTRACTING else None

            for run_idx, run in enumerate(paragraph.runs):
                text = run.text

                # If we found the right quote, accumulate remaining text for cite search
                if state == ExtractionState.FOUND_QUOTE:
                    cite_search_text += text

                if state == ExtractionState.SEARCHING:
                    # Look for bold left quote
                    if LEFT_QUOTE in text and run.bold:
                        logging.debug(f"Found LEFT_QUOTE at paragraph {para_idx}")
                        state = ExtractionState.EXTRACTING
                        para_html_start = len(accumulated_html)
                        # Calculate where extraction starts in paragraph text (after LEFT_QUOTE)
                        char_offset = sum(len(paragraph.runs[i].text) for i in range(run_idx))
                        para_extract_start = char_offset + text.index(LEFT_QUOTE) + 1
                        start_paragraph_index = para_idx

                        # Page number will be filled in by COM automation after extraction
                        page_number = None

                        # Split text at left quote and start accumulating
                        split_idx = text.index(LEFT_QUOTE)
                        text_after_quote = text[split_idx + 1:]

                        # Check if right quote is in the same run
                        if RIGHT_QUOTE in text_after_quote and run.bold:
                            # Complete translation in single run
                            end_idx = text_after_quote.index(RIGHT_QUOTE)
                            translation_text = text_after_quote[:end_idx]
                            cite_search_text = text_after_quote[end_idx + 1:]

                            # Format the extracted text with font awareness
                            formatted = replace_unicode_chars(translation_text)

                            # Check font
                            font_name = None
                            try:
                                if run.font and run.font.name and run.font.name != 'Times New Roman':
                                    font_name = run.font.name
                            except:
                                pass

                            if font_name:
                                formatted = f'<span class="{font_name}">{formatted}</span>'

                            if run.underline:
                                formatted = f"<u>{formatted}</u>"
                            if run.italic:
                                formatted = f"<i>{formatted}</i>"
                            if run.bold:
                                formatted = f"<b>{formatted}</b>"

                            accumulated_html.append(formatted)
                            state = ExtractionState.FOUND_QUOTE
                        else:
                            # Start accumulating
                            formatted = replace_unicode_chars(text_after_quote)

                            # Check font
                            font_name = None
                            try:
                                if run.font and run.font.name and run.font.name != 'Times New Roman':
                                    font_name = run.font.name
                            except:
                                pass

                            if font_name:
                                formatted = f'<span class="{font_name}">{formatted}</span>'

                            if run.underline:
                                formatted = f"<u>{formatted}</u>"
                            if run.italic:
                                formatted = f"<i>{formatted}</i>"
                            if run.bold:
                                formatted = f"<b>{formatted}</b>"

                            accumulated_html.append(formatted)

                elif state == ExtractionState.EXTRACTING:
                    # Continue accumulating until we find bold right quote
                    if RIGHT_QUOTE in text and run.bold:
                        logging.debug(f"Found RIGHT_QUOTE at paragraph {para_idx}")

                        # Split at right quote
                        split_idx = text.index(RIGHT_QUOTE)
                        text_before_quote = text[:split_idx]
                        cite_search_text = text[split_idx + 1:]

                        # Format the text before the quote
                        if text_before_quote:
                            formatted = replace_unicode_chars(text_before_quote)

                            # Check font
                            font_name = None
                            try:
                                if run.font and run.font.name and run.font.name != 'Times New Roman':
                                    font_name = run.font.name
                            except:
                                pass

                            if font_name:
                                formatted = f'<span class="{font_name}">{formatted}</span>'

                            if run.underline:
                                formatted = f"<u>{formatted}</u>"
                            if run.italic:
                                formatted = f"<i>{formatted}</i>"
                            if run.bold:
                                formatted = f"<b>{formatted}</b>"

                            accumulated_html.append(formatted)

                        state = ExtractionState.FOUND_QUOTE
                    else:
                        # Keep accumulating
                        formatted = format_run_as_html(run)
                        accumulated_html.append(formatted)

            # If still extracting at end of paragraph, check for cite-terminated translation
            if state == ExtractionState.EXTRACTING:
                para_text = replace_unicode_chars(paragraph.text)
                # Check if paragraph ends with balanced (...chapter:verse...) using paren matching
                cite_found = False
                stripped_pt = para_text.rstrip()
                if stripped_pt.endswith(')'):
                    depth = 0
                    paren_start = None
                    for i in range(len(stripped_pt) - 1, -1, -1):
                        if stripped_pt[i] == ')':
                            depth += 1
                        elif stripped_pt[i] == '(':
                            depth -= 1
                            if depth == 0:
                                paren_start = i
                                break
                    if paren_start is not None:
                        cite_content = stripped_pt[paren_start + 1:-1].strip()
                        if re.search(r'\d+:\d+', cite_content):
                            cite_found = True
                            raw_cite = cite_content
                            cite_start_pos = paren_start

                if cite_found:
                    extract_from = para_extract_start if para_extract_start is not None else 0

                    # Truncate accumulated_html back to before this paragraph's runs
                    if para_html_start is not None:
                        accumulated_html = accumulated_html[:para_html_start]

                    # Re-walk this paragraph's runs, building HTML for only the non-cite portion
                    char_pos = 0
                    for r in paragraph.runs:
                        r_text = replace_unicode_chars(r.text)
                        r_end = char_pos + len(r_text)

                        if r_end <= extract_from or char_pos >= cite_start_pos:
                            char_pos = r_end
                            continue

                        eff_start = max(0, extract_from - char_pos)
                        eff_end = min(len(r_text), cite_start_pos - char_pos)

                        if eff_end > eff_start:
                            keep = r_text[eff_start:eff_end]
                            if keep:
                                fmt = keep
                                fn = None
                                try:
                                    if r.font and r.font.name and r.font.name != 'Times New Roman':
                                        fn = r.font.name
                                except:
                                    pass
                                if fn:
                                    fmt = f'<span class="{fn}">{fmt}</span>'
                                if r.underline:
                                    fmt = f"<u>{fmt}</u>"
                                if r.italic:
                                    fmt = f"<i>{fmt}</i>"
                                if r.bold:
                                    fmt = f"<b>{fmt}</b>"
                                accumulated_html.append(fmt)

                        char_pos = r_end

                    cite_name, cite_chapter, cite_verse, cite_verse_end, cite_note = parse_cite(raw_cite)
                    cite_hebrew, cite_common = split_cite_name(cite_name)
                    full_text = consolidate_html("".join(accumulated_html))

                    if cite_name or cite_chapter:
                        translation = {
                            "book": book_name,
                            "page": None,
                            "text_word": full_text,
                            "cite": cite_name,
                            "cite_hebrew": cite_hebrew,
                            "cite_common": cite_common,
                            "cite_chapter": cite_chapter,
                            "cite_verse": cite_verse,
                            "cite_verse_end": cite_verse_end,
                            "cite_note": cite_note,
                            "_para_idx": start_paragraph_index
                        }
                        translations.append(translation)
                        logging.info(f"Extracted cite-terminated translation #{len(translations)} from {book_name}")

                    state = ExtractionState.SEARCHING
                    accumulated_html = []
                    cite_search_text = ""
                    page_number = None
                else:
                    accumulated_html.append("<br>")

            # Process completed extraction at end of paragraph
            if state == ExtractionState.FOUND_QUOTE:
                # Extract cite from accumulated cite search text
                raw_cite = extract_cite(cite_search_text)

                # Only save translations that have a cite reference
                if raw_cite:
                    cite_name, cite_chapter, cite_verse, cite_verse_end, cite_note = parse_cite(raw_cite)
                    cite_hebrew, cite_common = split_cite_name(cite_name)

                    # Combine accumulated HTML and consolidate adjacent tags
                    full_text = consolidate_html("".join(accumulated_html))

                    translation = {
                        "book": book_name,
                        "page": None,
                        "text_word": full_text,
                        "cite": cite_name,
                        "cite_hebrew": cite_hebrew,
                        "cite_common": cite_common,
                        "cite_chapter": cite_chapter,
                        "cite_verse": cite_verse,
                        "cite_verse_end": cite_verse_end,
                        "cite_note": cite_note,
                        "_para_idx": start_paragraph_index
                    }

                    translations.append(translation)
                    logging.info(f"Extracted translation #{len(translations)} from {book_name}")
                    logging.debug(f"  Cite: {cite_name} {cite_chapter}:{cite_verse}")
                    logging.debug(f"  Text preview: {full_text[:100]}...")
                else:
                    logging.debug(f"Skipped translation without cite at paragraph {start_paragraph_index}")

                # Reset state machine
                state = ExtractionState.SEARCHING
                accumulated_html = []
                cite_search_text = ""
                page_number = None

        # Handle case where we finished the document while still searching for cite
        if state == ExtractionState.FOUND_QUOTE and len(accumulated_html) > 0:
            raw_cite = extract_cite(cite_search_text)
            if raw_cite:
                cite_name, cite_chapter, cite_verse, cite_verse_end, cite_note = parse_cite(raw_cite)
                cite_hebrew, cite_common = split_cite_name(cite_name)
                full_text = consolidate_html("".join(accumulated_html))

                translation = {
                    "book": book_name,
                    "page": None,
                    "text_word": full_text,
                    "cite": cite_name,
                    "cite_hebrew": cite_hebrew,
                    "cite_common": cite_common,
                    "cite_chapter": cite_chapter,
                    "cite_verse": cite_verse,
                    "cite_verse_end": cite_verse_end,
                    "cite_note": cite_note,
                    "_para_idx": start_paragraph_index
                }
                translations.append(translation)

        if state == ExtractionState.EXTRACTING:
            logging.warning(f"Document {book_name} has unclosed translation (missing right quote)")

        return translations

    except Exception as e:
        logging.error(f"Error processing document {doc_path}: {e}")
        return []


def save_translation(conn, translation_data: Dict[str, any]) -> bool:
    """
    Insert translation record into database.

    Args:
        conn: PostgreSQL connection object
        translation_data: Dictionary with keys: book, page, text_word, cite, cite_chapter, cite_verse

    Returns:
        True if successful, False otherwise
    """
    try:
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO translation (translation_book, translation_page, translation_text_word,
                                     translation_cite, translation_cite_hebrew, translation_cite_common,
                                     translation_cite_chapter, translation_cite_verse,
                                     translation_cite_verse_end, translation_cite_note)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            translation_data['book'],
            translation_data['page'],
            translation_data['text_word'],
            translation_data['cite'],
            translation_data.get('cite_hebrew'),
            translation_data.get('cite_common'),
            translation_data.get('cite_chapter'),
            translation_data.get('cite_verse'),
            translation_data.get('cite_verse_end'),
            translation_data.get('cite_note')
        ))

        conn.commit()
        cursor.close()
        return True

    except psycopg2.Error as e:
        logging.error(f"Database insertion failed: {e}")
        conn.rollback()
        return False


def populate_cite_table(conn) -> int:
    """
    Populate the cite table with distinct cite values from the translation table.

    Returns:
        Number of distinct cite values inserted
    """
    try:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO cite (label)
            SELECT DISTINCT translation_cite FROM translation
            WHERE translation_cite IS NOT NULL
            ON CONFLICT (label) DO NOTHING
        """)
        count = cursor.rowcount
        conn.commit()
        cursor.close()
        logging.info(f"Populated cite table with {count} distinct cite values")
        return count
    except psycopg2.Error as e:
        logging.error(f"Failed to populate cite table: {e}")
        conn.rollback()
        return 0


def update_cite_book_ids(conn) -> int:
    """
    Update translation.cite_book_id by looking up cite_hebrew in cite_book_map.

    Returns:
        Number of rows updated
    """
    try:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE translation t
            SET translation_cite_book_id = cbm.cite_book_id
            FROM cite_book_map cbm
            WHERE t.translation_cite_hebrew = cbm.cite_book_map_hebrew
              AND t.translation_cite_book_id IS NULL
        """)
        count = cursor.rowcount
        conn.commit()
        cursor.close()
        logging.info(f"Updated cite_book_id for {count} translation rows")
        return count
    except psycopg2.Error as e:
        logging.error(f"Failed to update cite_book_id: {e}")
        conn.rollback()
        return 0


def parse_directory(directory_path: Path, conn, dry_run: bool = False) -> Dict[str, int]:
    """
    Parse all Word documents in a directory.

    Args:
        directory_path: Path to directory containing .docx files
        conn: PostgreSQL connection object
        dry_run: If True, extract but don't save to database

    Returns:
        Dictionary with statistics: files_processed, translations_found, translations_saved
    """
    stats = {
        "files_processed": 0,
        "translations_found": 0,
        "translations_saved": 0
    }

    if not directory_path.exists():
        logging.error(f"Directory does not exist: {directory_path}")
        return stats

    # Find all .docx files starting with "YY" (excluding temporary files starting with ~$)
    docx_files = [f for f in directory_path.glob("YY*.docx")
                  if not f.name.startswith("~$") and not f.name.startswith("YY-s07")]

    if not docx_files:
        logging.warning(f"No .docx files found in {directory_path}")
        return stats

    logging.info(f"Found {len(docx_files)} document(s) to process")

    # Phase 1: Extract all translations using python-docx (fast)
    all_translations = {}  # doc_path -> list of translations
    for doc_path in docx_files:
        logging.info(f"\n{'='*60}")
        translations = extract_translations_from_doc(doc_path)
        all_translations[doc_path] = translations
        stats["files_processed"] += 1
        stats["translations_found"] += len(translations)

    # Phase 2: Get page numbers via single COM subprocess
    doc_para_map = {}
    for doc_path, translations in all_translations.items():
        if translations:
            para_indices = list(set(t["_para_idx"] for t in translations))
            doc_para_map[str(doc_path.absolute())] = para_indices

    if doc_para_map and not dry_run:
        all_page_numbers = get_all_page_numbers(doc_para_map)

        # Apply page numbers to translations
        for doc_path, translations in all_translations.items():
            abs_path = str(doc_path.absolute())
            page_map = all_page_numbers.get(abs_path, {})
            for t in translations:
                pidx = t.pop("_para_idx")
                t["page"] = page_map.get(pidx)
    else:
        # Remove _para_idx from translations even if not using COM
        for doc_path, translations in all_translations.items():
            for t in translations:
                t.pop("_para_idx", None)

    # Phase 3: Save to database
    for doc_path, translations in all_translations.items():
        if dry_run:
            logging.info(f"[DRY RUN] Would save {len(translations)} translation(s) from {doc_path.name}")
            for i, trans in enumerate(translations, 1):
                logging.info(f"  Translation {i}:")
                logging.info(f"    Book: {trans['book']}")
                logging.info(f"    Page: {trans.get('page')}")
                logging.info(f"    Cite: {trans['cite']}")
                logging.info(f"    Text: {trans['text_word'][:100]}...")
        else:
            for translation in translations:
                if save_translation(conn, translation):
                    stats["translations_saved"] += 1

    return stats


def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description="Extract translations from Word documents to PostgreSQL database"
    )
    parser.add_argument(
        "--directory",
        type=str,
        default=DEFAULT_DIRECTORY,
        help=f"Directory containing .docx files (default: {DEFAULT_DIRECTORY})"
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Preview extractions without saving to database"
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable verbose logging"
    )

    args = parser.parse_args()

    setup_logging(args.verbose)

    logging.info("Word Document Translation Extraction Script")
    logging.info("=" * 60)

    # Validate directory
    directory_path = Path(args.directory)
    if not directory_path.exists():
        logging.error(f"Directory does not exist: {directory_path}")
        sys.exit(1)

    logging.info(f"Target directory: {directory_path}")
    logging.info(f"Dry run mode: {args.dry_run}")

    # Connect to database
    conn = None
    if not args.dry_run:
        conn = get_db_connection()
        conn = init_database(conn)
    else:
        logging.info("Skipping database connection (dry run mode)")

    # Process documents
    try:
        stats = parse_directory(directory_path, conn, dry_run=args.dry_run)

        logging.info("\n" + "=" * 60)
        logging.info("SUMMARY")
        logging.info("=" * 60)
        # Populate cite table with distinct cite values
        if not args.dry_run and conn:
            populate_cite_table(conn)
            update_cite_book_ids(conn)

        logging.info(f"Files processed: {stats['files_processed']}")
        logging.info(f"Translations found: {stats['translations_found']}")
        if not args.dry_run:
            logging.info(f"Translations saved: {stats['translations_saved']}")

        if stats['translations_found'] > 0 and not args.dry_run:
            if stats['translations_saved'] == stats['translations_found']:
                logging.info("✓ All translations saved successfully!")
            else:
                logging.warning(f"⚠ {stats['translations_found'] - stats['translations_saved']} translation(s) failed to save")

    finally:
        if conn:
            conn.close()
            logging.info("Database connection closed")


if __name__ == "__main__":
    main()
